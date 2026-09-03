"""Layout-aware local document parsing for grounded AI extraction.

PDF strategy:
- PyMuPDF remains the canonical text source and fallback parser.
- OpenDataLoader PDF is used when available for table structure / reading order.
- PyMuPDF ``find_tables`` is a deterministic fallback when OpenDataLoader or Java
  is unavailable.
- scoring sections emitted as ``Tip: OPTIUNI`` are parsed deterministically so
  only the option explicitly marked ``Selectată: Da`` reaches obligation
  extraction. These sections are often visually list-like, not actual PDF
  tables, so a table parser alone cannot solve them.

No cloud OCR or LLM is used here. The extracted wording remains local Romanian
source text; only whitespace and table separators are normalized.
"""

from __future__ import annotations

import io
import json
import os
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from uuid import UUID


@dataclass(frozen=True, slots=True)
class StructuredBlock:
    # ``text`` is the semantic representation shown to retrieval/Qwen. For a
    # table it may contain synthetic header/value separators that never appear
    # verbatim in the PDF. ``source_text`` is therefore kept separately and is
    # always a contiguous substring of the canonical page text.
    text: str
    kind: str  # table_row | selected_option
    source_text: str | None = None
    source_page_number: int | None = None


@dataclass(frozen=True, slots=True)
class ParsedPage:
    document_id: UUID
    page_number: int
    text: str
    blocks: tuple[StructuredBlock, ...] = ()
    # Scoring-option pages must not also be chunked as raw text, otherwise
    # unselected alternatives ("Selectată: Nu") can re-enter the LLM path.
    prefer_structured: bool = False


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    document_id: UUID
    pages: tuple[ParsedPage, ...]
    warnings: tuple[str, ...] = ()


def _clean_cell_text(value: Any) -> str:
    return " ".join(str(value or "").replace("\u00a0", " ").split()).strip()


def _normalized_whitespace_map(text: str) -> tuple[str, list[int]]:
    """Collapse whitespace while retaining a mapping back to source offsets."""
    normalized: list[str] = []
    source_offsets: list[int] = []
    in_space = False
    for index, char in enumerate(text):
        if char.isspace() or char == "\u00a0":
            if normalized and not in_space:
                normalized.append(" ")
                source_offsets.append(index)
            in_space = True
            continue
        normalized.append(char)
        source_offsets.append(index)
        in_space = False
    return "".join(normalized), source_offsets


def _source_span_from_fragments(page_text: str, fragments: list[str]) -> str | None:
    """Recover one exact contiguous page substring for a structured row.

    Table parsers are allowed to build a richer semantic representation, but
    persisted provenance must come from the original PyMuPDF page text. The
    row is anchored on its rarest/most distinctive cell, then the nearest
    occurrences of the remaining cells are selected. This avoids accidentally
    spanning several rows when generic values such as ``REPER`` repeat.
    """
    normalized_page, offsets = _normalized_whitespace_map(page_text)
    if not normalized_page or not offsets:
        return None

    needles: list[str] = []
    for fragment in fragments:
        cleaned = _clean_cell_text(fragment)
        if len(cleaned) < 3 or cleaned in {"-", "–", "—"}:
            continue
        needle = " ".join(cleaned.split())
        if needle not in needles:
            needles.append(needle)
    if not needles:
        return None

    occurrences: dict[str, list[tuple[int, int]]] = {}
    for needle in needles:
        matches: list[tuple[int, int]] = []
        cursor = 0
        while True:
            found = normalized_page.find(needle, cursor)
            if found < 0:
                break
            matches.append((found, found + len(needle)))
            cursor = found + max(1, len(needle))
        if matches:
            occurrences[needle] = matches
    if not occurrences:
        return None

    anchor_needle = min(
        occurrences,
        key=lambda needle: (len(occurrences[needle]), -len(needle)),
    )
    anchor = occurrences[anchor_needle][0]
    anchor_mid = (anchor[0] + anchor[1]) // 2
    positions = [anchor]
    for needle, matches in occurrences.items():
        if needle == anchor_needle:
            continue
        nearest = min(
            matches,
            key=lambda item: abs(((item[0] + item[1]) // 2) - anchor_mid),
        )
        # Do not pull in a same-looking cell from a distant table/section.
        if abs(((nearest[0] + nearest[1]) // 2) - anchor_mid) <= 2400:
            positions.append(nearest)

    first = min(start for start, _ in positions)
    last = max(end for _, end in positions)
    source_start = offsets[first]
    source_end = offsets[last - 1] + 1
    candidate = page_text[source_start:source_end].strip()

    # Avoid turning a difficult table mapping into an entire-page description.
    # A single distinctive exact cell is preferable to synthetic provenance.
    if len(candidate) > 3800:
        start, end = anchor
        source_start = offsets[start]
        source_end = offsets[end - 1] + 1
        candidate = page_text[source_start:source_end].strip()
    return candidate or None


def _row_text(headers: list[str], values: list[str]) -> str:
    """Serialize one table row without rewriting any cell wording."""
    values = [_clean_cell_text(value) for value in values]
    if not any(values):
        return ""
    if headers and len(headers) == len(values):
        parts = []
        for header, value in zip(headers, values, strict=False):
            if not value:
                continue
            header = _clean_cell_text(header)
            parts.append(f"{header}: {value}" if header else value)
        return " | ".join(parts)
    return " | ".join(value for value in values if value)


def _looks_like_header(values: list[str]) -> bool:
    if not values or not any(values):
        return False
    joined = " ".join(values)
    # Headers are usually mostly words and short; a row containing dates / long
    # values should not be consumed as a header.
    return bool(re.search(r"[A-Za-zĂÂÎȘȚăâîșț]{3,}", joined)) and len(joined) < 700


def _pymupdf_table_blocks(pdf: Any) -> dict[int, list[StructuredBlock]]:
    """Extract structured rows with PyMuPDF when ODL is unavailable."""
    by_page: dict[int, list[StructuredBlock]] = {}
    for page_number, page in enumerate(pdf, start=1):
        page_text = page.get_text("text") or ""
        try:
            found = page.find_tables()
        except Exception:
            continue
        for table in getattr(found, "tables", ()):
            try:
                rows = table.extract()
            except Exception:
                continue
            cleaned = [
                [_clean_cell_text(cell) for cell in row]
                for row in rows
                if row and any(_clean_cell_text(cell) for cell in row)
            ]
            if not cleaned:
                continue
            headers: list[str] = []
            start = 0
            if len(cleaned) > 1 and _looks_like_header(cleaned[0]):
                headers = cleaned[0]
                start = 1
            for row in cleaned[start:]:
                text = _row_text(headers, row)
                if len(text) >= 12:
                    by_page.setdefault(page_number, []).append(
                        StructuredBlock(
                            text=text,
                            kind="table_row",
                            source_text=_source_span_from_fragments(page_text, row),
                            source_page_number=page_number,
                        )
                    )
    return by_page


def _node_text(node: Any) -> str:
    if not isinstance(node, dict):
        return ""
    pieces: list[str] = []
    content = node.get("content")
    if isinstance(content, str) and content.strip():
        pieces.append(content)
    for key in ("kids", "list items"):
        children = node.get(key)
        if isinstance(children, list):
            for child in children:
                text = _node_text(child)
                if text:
                    pieces.append(text)
    return _clean_cell_text(" ".join(pieces))


def _walk_nodes(node: Any):
    if isinstance(node, dict):
        yield node
        for value in node.values():
            if isinstance(value, (dict, list)):
                yield from _walk_nodes(value)
    elif isinstance(node, list):
        for item in node:
            yield from _walk_nodes(item)


def _opendataloader_table_blocks(
    content: bytes, page_text_by_number: dict[int, str]
) -> dict[int, list[StructuredBlock]]:
    """Use OpenDataLoader PDF JSON to preserve table row/column relationships.

    OpenDataLoader is deliberately optional at runtime.  The project still works
    with the PyMuPDF fallback when Java/package installation is unavailable.
    """
    if os.getenv("CHIATRATON_PDF_TABLE_BACKEND", "auto").lower() == "pymupdf":
        raise RuntimeError("OpenDataLoader disabled by configuration")
    if shutil.which("java") is None:
        raise RuntimeError("Java is not available")
    try:
        import opendataloader_pdf
    except ImportError as exc:
        raise RuntimeError("opendataloader-pdf is not installed") from exc

    with tempfile.TemporaryDirectory(prefix="chiatraton-odl-") as directory:
        root = Path(directory)
        input_path = root / "document.pdf"
        output_dir = root / "out"
        output_dir.mkdir()
        input_path.write_bytes(content)
        opendataloader_pdf.convert(
            input_path=str(input_path),
            output_dir=str(output_dir),
            format="json",
            table_method="cluster",
            reading_order="xycut",
            quiet=True,
        )
        json_files = sorted(output_dir.rglob("*.json"))
        if not json_files:
            raise RuntimeError("OpenDataLoader produced no JSON output")
        data = json.loads(json_files[0].read_text(encoding="utf-8"))

    by_page: dict[int, list[StructuredBlock]] = {}
    for node in _walk_nodes(data.get("kids", [])):
        if node.get("type") != "table":
            continue
        rows = node.get("rows")
        if not isinstance(rows, list) or not rows:
            continue
        page_number = int(node.get("page number") or 1)
        extracted_rows: list[list[str]] = []
        for row in rows:
            if not isinstance(row, dict):
                continue
            cells = row.get("cells")
            if not isinstance(cells, list):
                continue
            cells = sorted(
                (cell for cell in cells if isinstance(cell, dict)),
                key=lambda cell: int(cell.get("column number") or 0),
            )
            extracted_rows.append([_node_text(cell) for cell in cells])
        if not extracted_rows:
            continue
        headers: list[str] = []
        start = 0
        if len(extracted_rows) > 1 and _looks_like_header(extracted_rows[0]):
            headers = extracted_rows[0]
            start = 1
        for row in extracted_rows[start:]:
            text = _row_text(headers, row)
            if len(text) >= 12:
                by_page.setdefault(page_number, []).append(
                    StructuredBlock(
                        text=text,
                        kind="table_row",
                        source_text=_source_span_from_fragments(
                            page_text_by_number.get(page_number, ""), row
                        ),
                        source_page_number=page_number,
                    )
                )
    return by_page


_OPTION_RE = re.compile(
    r"Descriere:\s*(?P<description>.*?)"
    r";\s*(?P<score>\d+(?:[.,]\d+)?)\s*;\s*"
    r"(?P<selected>Da|Nu)\s*Punctaj:\s*Selectat[ăa]:",
    re.IGNORECASE | re.DOTALL,
)
_SUBCRITERION_RE = re.compile(
    r"Descriere subcriteriu:\s*(?P<context>.*?)\s*Tip:\s*OPTIUNI",
    re.IGNORECASE | re.DOTALL,
)
_HISTORICAL_SCORING = re.compile(
    r"(rata\s+solvabilit|anul\s+fiscal\s+anterior|"
    r"raportul\s+dintre\s+cuantumul\s+finanțării\s+nerambursabile.*cifra\s+de\s+afaceri|"
    r"\bsold\s+(?:negativ|pozitiv))",
    re.IGNORECASE | re.DOTALL,
)
_MONITORABLE_SCORING = re.compile(
    r"(se\s+angajeaz|angajarea|locuri?\s+de\s+muncă|salariaț|defavorizat|"
    r"mențin|monitorizare|până\s+la|contribuția\s+solicitantului|"
    r"localizat|localizare|județ|sediul\s+social|"
    r"investiția\s+(?:prevede|propune)|durabil|emisii|eficienț[aă]\s+din\s+punct\s+de\s+vedere)",
    re.IGNORECASE,
)


def _selected_option_blocks(page_texts: list[str]) -> tuple[dict[int, list[StructuredBlock]], set[int]]:
    """Parse MySMIS scoring alternatives and keep only explicit selections.

    Context is carried across page boundaries because long option lists (e.g.
    contribution 11%-20%) continue on the next page. Semantic text may be
    normalized for the model, while ``source_text`` preserves a verbatim
    contiguous substring from the originating PDF page.
    """
    by_page: dict[int, list[StructuredBlock]] = {}
    option_pages: set[int] = set()
    current_context = ""
    current_context_raw = ""
    current_context_page: int | None = None

    for page_number, text in enumerate(page_texts, start=1):
        contexts = list(_SUBCRITERION_RE.finditer(text))
        if contexts:
            last = contexts[-1]
            current_context = _clean_cell_text(last.group("context"))
            current_context_raw = last.group("context").strip()
            current_context_page = page_number
        has_option_structure = bool(contexts) or bool(_OPTION_RE.search(text))
        if has_option_structure:
            option_pages.add(page_number)

        for match in _OPTION_RE.finditer(text):
            # Find the closest subcriterion context before this option on the
            # same page; otherwise inherit the prior-page context.
            local_context = current_context
            local_context_raw = current_context_raw
            local_context_page = current_context_page
            prior = [ctx for ctx in contexts if ctx.start() < match.start()]
            if prior:
                context_match = prior[-1]
                local_context = _clean_cell_text(context_match.group("context"))
                local_context_raw = context_match.group("context").strip()
                local_context_page = page_number
                current_context = local_context
                current_context_raw = local_context_raw
                current_context_page = local_context_page
            if match.group("selected").lower() != "da":
                continue
            try:
                score_value = float(match.group("score").replace(",", "."))
            except ValueError:
                score_value = 0.0
            if score_value <= 0:
                continue

            description_raw = match.group("description").strip()
            description = _clean_cell_text(description_raw)
            context = _clean_cell_text(local_context)
            combined_context = f"{context} {description}".strip()
            # Application-time financial measurements are evaluation facts, not
            # future obligations.  Do this deterministically rather than asking
            # the LLM to rediscover it each run.
            if _HISTORICAL_SCORING.search(combined_context):
                continue
            if not _MONITORABLE_SCORING.search(combined_context):
                continue

            # For binary Da/Nu options the useful source wording is the
            # subcriterion itself, not the isolated word "Da".
            if re.fullmatch(r"(?:[a-z]\.?\s*)?Da", description, re.IGNORECASE):
                selected_text = context
                source_text = local_context_raw
                source_page_number = local_context_page or page_number
            else:
                selected_text = description
                source_text = description_raw
                source_page_number = page_number
            if not selected_text or not source_text:
                continue

            # Selection metadata is useful semantic context for Qwen but is not
            # persisted as the obligation passage.
            block_text = (
                f"{selected_text}\n"
                f"Punctaj: {_clean_cell_text(match.group('score'))}; Selectată: Da"
            )
            by_page.setdefault(page_number, []).append(
                StructuredBlock(
                    text=block_text,
                    kind="selected_option",
                    source_text=source_text,
                    source_page_number=source_page_number,
                )
            )

    return by_page, option_pages


def _pdf(document_id: UUID, content: bytes) -> ParsedDocument:
    try:
        import pymupdf
    except ImportError as exc:  # pragma: no cover - depends on optional AI extra
        raise RuntimeError("Install the 'ai' extra: pymupdf is required for PDF parsing") from exc

    base_pages: list[tuple[int, str]] = []
    empty: list[int] = []
    with pymupdf.open(stream=content, filetype="pdf") as pdf:
        for page_number, page in enumerate(pdf, start=1):
            text = page.get_text("text") or ""
            if len(text.strip()) < 10:
                empty.append(page_number)
                continue
            base_pages.append((page_number, text))

        table_backend = "pymupdf"
        table_blocks: dict[int, list[StructuredBlock]] = {}
        backend = os.getenv("CHIATRATON_PDF_TABLE_BACKEND", "auto").lower()
        if backend not in {"auto", "opendataloader", "pymupdf"}:
            backend = "auto"
        if backend in {"auto", "opendataloader"}:
            try:
                table_blocks = _opendataloader_table_blocks(
                    content, {number: text for number, text in base_pages}
                )
                table_backend = "opendataloader"
            except Exception:
                if backend == "opendataloader":
                    # Explicit ODL mode still falls back instead of breaking a
                    # hackathon/demo workflow because the evidence parser must
                    # remain available offline.
                    table_backend = "pymupdf-fallback"
        if not table_blocks:
            table_blocks = _pymupdf_table_blocks(pdf)
            if table_backend == "opendataloader":
                table_backend = "opendataloader+pymupdf-fallback"

    text_by_number = {number: text for number, text in base_pages}
    max_page = max(text_by_number, default=0)
    page_texts = [text_by_number.get(number, "") for number in range(1, max_page + 1)]
    option_blocks, option_pages = _selected_option_blocks(page_texts)

    pages: list[ParsedPage] = []
    for page_number, text in base_pages:
        blocks = tuple(
            table_blocks.get(page_number, []) + option_blocks.get(page_number, [])
        )
        pages.append(
            ParsedPage(
                document_id=document_id,
                page_number=page_number,
                text=text,
                blocks=blocks,
                prefer_structured=page_number in option_pages,
            )
        )

    warnings: list[str] = [f"PDF structured-table backend: {table_backend}"]
    if empty:
        preview = ", ".join(map(str, empty[:12])) + ("..." if len(empty) > 12 else "")
        warnings.append(
            f"{len(empty)} PDF page(s) have no extractable text layer: {preview}; "
            "OCR was not performed because evidence must preserve source wording."
        )
    return ParsedDocument(document_id, tuple(pages), tuple(warnings))


def _docx(document_id: UUID, content: bytes) -> ParsedDocument:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("Install the 'ai' extra: python-docx is required for DOCX") from exc

    doc = DocxDocument(io.BytesIO(content))
    blocks: list[str] = []
    pages: list[ParsedPage] = []
    page_number = 1

    def flush() -> None:
        nonlocal blocks, page_number
        text = "\n".join(blocks).strip()
        if text:
            pages.append(ParsedPage(document_id, page_number, text))
            page_number += 1
        blocks = []

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not text.strip():
            continue
        blocks.append(text)
        if sum(len(item) for item in blocks) >= 5000:
            flush()
    flush()
    return ParsedDocument(document_id, tuple(pages))


def _xlsx(document_id: UUID, content: bytes) -> ParsedDocument:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("Install the 'ai' extra: openpyxl is required for XLSX") from exc

    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    pages: list[ParsedPage] = []
    page_number = 1
    for sheet in workbook.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                rows.append(" | ".join(values))
            if sum(len(item) for item in rows) >= 4500:
                text = "\n".join(rows).strip()
                if text:
                    pages.append(ParsedPage(document_id, page_number, text))
                    page_number += 1
                rows = []
        if rows:
            pages.append(ParsedPage(document_id, page_number, "\n".join(rows).strip()))
            page_number += 1
    return ParsedDocument(document_id, tuple(pages))


def _xls(document_id: UUID, content: bytes) -> ParsedDocument:
    try:
        import xlrd
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("Install the 'ai' extra: xlrd is required for legacy XLS") from exc

    workbook = xlrd.open_workbook(file_contents=content)
    pages: list[ParsedPage] = []
    page_number = 1
    for sheet in workbook.sheets():
        rows: list[str] = []
        for row_index in range(sheet.nrows):
            values = [str(value).strip() for value in sheet.row_values(row_index) if str(value).strip()]
            if values:
                rows.append(" | ".join(values))
            if sum(len(item) for item in rows) >= 4500:
                pages.append(ParsedPage(document_id, page_number, "\n".join(rows).strip()))
                page_number += 1
                rows = []
        if rows:
            pages.append(ParsedPage(document_id, page_number, "\n".join(rows).strip()))
            page_number += 1
    return ParsedDocument(document_id, tuple(pages))


def _doc(document_id: UUID, content: bytes) -> ParsedDocument:
    antiword = shutil.which("antiword")
    if not antiword:
        raise RuntimeError("Legacy DOC requires local 'antiword' or conversion to DOCX")
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as handle:
        path = Path(handle.name)
        handle.write(content)
    try:
        result = subprocess.run(
            [antiword, str(path)],
            capture_output=True,
            text=True,
            check=True,
        )
    finally:
        path.unlink(missing_ok=True)
    return ParsedDocument(document_id, (ParsedPage(document_id, 1, result.stdout),))


def parse_document_bytes(document_id: UUID, media_type: str, content: bytes) -> ParsedDocument:
    if not content:
        raise RuntimeError("Document content is empty")
    if media_type == "application/pdf":
        return _pdf(document_id, content)
    if media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _docx(document_id, content)
    if media_type == "application/msword":
        return _doc(document_id, content)
    if media_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        return _xlsx(document_id, content)
    if media_type == "application/vnd.ms-excel":
        return _xls(document_id, content)
    raise RuntimeError(f"Unsupported AI document media type: {media_type}")
